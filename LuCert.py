from qtUI import Ui_MainWindow
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook, load_workbook
from PyQt5 import QtWidgets, QtCore
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QMessageBox, QErrorMessage
from PyQt5.QtCore import QDir
from docx2pdf import convert
import openpyxl
import platform
import docx
import sys
import os
import shutil
import itertools as it



class MyWindow(QMainWindow, Ui_MainWindow):

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.setWindowTitle("LuCert")

        self.setupUi(self)

        global emsg
        emsg = QtWidgets.QErrorMessage(self)
        emsg.setWindowModality(QtCore.Qt.WindowModal)

        self.rdbSep.toggled.connect(self.onClicked)
        self.btnExc.clicked.connect(self.exc_clicked)
        self.btnUploadXl.clicked.connect(self.upload_clicked)
        self.btnUploadDocx.clicked.connect(self.uploadDocx_clicked)
        self.cmbFN.setEnabled(False)
        self.cmbLN.setEnabled(False)

    def onClicked(self):
        rdbSep = self.sender()
        if rdbSep.isChecked():
            self.cmbFN.setEnabled(True)
            self.cmbLN.setEnabled(True)

    def exc_clicked(self):
        self.update()

        cmbCol_selected = self.cmbCol.currentText()
        print("1 "+os.getcwd())
        os.chdir(os.path.abspath('Lists'))
        print("2 "+os.getcwd())

        wb = load_workbook("List.xlsx")
        source = wb.worksheets[0]

        if self.rdbSep.isChecked():
            print("3 "+os.getcwd())
            cmbFN_selected = self.cmbFN.currentText()
            cmbLN_selected = self.cmbLN.currentText()

            for cell1,cell2 in zip(source[cmbFN_selected],source[cmbLN_selected]):

                name = cell1.value+" "+cell2.value
                if cell1.value is None or cell2.value is None:
                    emsg.setWindowTitle("Cell contains nothing")
                    emsg.showMessage("The coloumn contains nothing, please correct in xlsx file!")

                os.chdir(os.path.abspath('..'))

                print("4 "+os.getcwd())
                doc = docx.Document('masterDoc.docx')
                p1 = doc.add_paragraph(name)
                p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if os.path.exists(os.path.abspath("Certs")):
                    print("5 "+os.getcwd())
                    os.chdir(os.path.abspath("Certs"))

                    doc.save('%s.docx' % name)

                    if platform.system() == 'Windows':
                        convert('%s.docx' % name,'%s.pdf' % name)
                    else:
                        print("Not windows")

                    print("6 "+os.getcwd())
                else:
                    print('failed')

                wb.save("List.xlsx")

            os.chdir(os.path.abspath('..'))
            print("7 "+os.getcwd())

        else:

            for cell in source[cmbCol_selected]:
                print("8 "+os.getcwd())
                name =cell.value
                print(cell.value)

                if name is None:
                    emsg.setWindowTitle("Cell contains nothing")
                    emsg.showMessage("The coloumn contains nothing, please correct in xlsx file!")

                os.chdir(os.path.abspath('..'))

                print("9 "+os.getcwd())
                doc = docx.Document('masterDoc.docx')
                p1 = doc.add_paragraph(name)
                p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if os.path.exists(os.path.abspath("Certs")):
                    print("10 "+os.getcwd())

                    os.chdir(os.path.abspath("Certs"))

                    doc.save('%s.docx' % cell.value)

                    if platform.system() == 'Windows':
                        convert('%s.docx' % name,'%s.pdf' % name)
                    else:
                        print("Not windows")

                    print("11 "+os.getcwd())
                else:
                    print('failed')

            os.chdir(os.path.abspath('..'))

            print("12 "+os.getcwd())

    def uploadDocx_clicked(self):
        print("13 "+os.getcwd())
        docxDir = QFileDialog.getOpenFileName(self)
        docxDir = docxDir[0]
        self.txeDocx.setText(docxDir)
        shutil.copyfile(docxDir,os.getcwd()+'masterDoc.docx')
        print("14 "+os.getcwd())

        if docxDir is None:
            emsg.setWindowTitle("No Input")
            emsg.showMessage("The chosen file contains nothing")

        if docxDir == '':
            emsg.setWindowTitle("No Input")
            emsg.showMessage("The chosen file contains nothing")

        if not docxDir.endswith('.docx'):
            #print("ERROR")
            emsg.setWindowTitle("Wrong file")
            emsg.showMessage("The chosen file is not in .docx format!")

    def upload_clicked(self):
        print("15 "+os.getcwd())
        xlsxListDir = QFileDialog.getOpenFileName(self)
        xlsxListDir = xlsxListDir[0]
        self.txeXl.setText(xlsxListDir)
        shutil.copyfile(xlsxListDir,os.getcwd()+'/Lists/List.xlsx')
        print("16 "+os.getcwd())

        if xlsxListDir is None:
            emsg.setWindowTitle("No Input")
            emsg.showMessage("The chosen file contains nothing")

        if xlsxListDir == '':
            emsg.setWindowTitle("No Input")
            emsg.showMessage("The chosen file contains nothing")

        if not xlsxListDir.endswith('.xlsx'):
            emsg.setWindowTitle("Wrong file")
            emsg.showMessage("The chosen file is not in .xlsx format!")


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = MyWindow()
    #MainWindow = QtWidgets.QMainWindow()
    # ui = Ui_MainWindow()
    # ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
